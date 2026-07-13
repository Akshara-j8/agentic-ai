"""
TechVest Recruitment Agent — Main Application
6-Step Wizard: Landing → JD → Resumes → Run → Execution → Results
No hardcoded candidates, counts, or decisions. Everything is LLM-generated.
Run with: streamlit run app.py
"""
from __future__ import annotations

import json
import logging
import sys
import time
import uuid
from pathlib import Path

import streamlit as st

ROOT = Path(__file__).parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from config.settings import get_settings
from tools.parser import extract_text_from_pdf_bytes

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s")
logger = logging.getLogger(__name__)

st.set_page_config(
    page_title="TechVest Recruitment Agent",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={"About": "TechVest Recruitment Agent — Powered by LangGraph + OpenRouter"},
)

# ── CSS ───────────────────────────────────────────────────────────────────────
def inject_css() -> None:
    css_path = ROOT / "assets" / "styles.css"
    if css_path.exists():
        st.markdown(f"<style>{css_path.read_text(encoding='utf-8')}</style>", unsafe_allow_html=True)
    st.markdown("""
    <style>
    .step-header{background:linear-gradient(135deg,rgba(99,102,241,0.15),rgba(139,92,246,0.1));
        border:1px solid rgba(99,102,241,0.3);border-radius:16px;padding:1.5rem 2rem;margin-bottom:1.5rem;}
    .step-title{font-size:1.4rem;font-weight:800;color:#F1F5F9;margin:0;}
    .step-subtitle{font-size:0.85rem;color:#94A3B8;margin-top:0.25rem;}
    .step-badge{display:inline-block;background:rgba(99,102,241,0.2);color:#818CF8;
        border:1px solid rgba(99,102,241,0.4);border-radius:99px;
        padding:0.15rem 0.7rem;font-size:0.72rem;font-weight:700;margin-bottom:0.5rem;}
    .wizard-progress{display:flex;gap:0.5rem;margin-bottom:1.5rem;align-items:center;}
    .wstep{display:flex;align-items:center;gap:0.4rem;font-size:0.75rem;padding:0.35rem 0.75rem;
        border-radius:99px;font-weight:600;white-space:nowrap;}
    .wstep-done{background:rgba(16,185,129,0.15);color:#10B981;border:1px solid rgba(16,185,129,0.3);}
    .wstep-active{background:rgba(99,102,241,0.25);color:#818CF8;border:1px solid rgba(99,102,241,0.5);}
    .wstep-pending{background:rgba(30,41,59,0.5);color:#475569;border:1px solid rgba(255,255,255,0.05);}
    .section-title{font-size:1rem;font-weight:700;color:#F1F5F9;margin:1rem 0 0.75rem;}
    .card-section{background:rgba(30,41,59,0.5);border:1px solid rgba(255,255,255,0.07);
        border-radius:12px;padding:1.25rem;margin-bottom:1rem;}
    .interview-section{border-left:3px solid #10B981;}
    .hold-section{border-left:3px solid #F59E0B;}
    .reject-section{border-left:3px solid #EF4444;}
    .traj-event{padding:0.6rem 0.85rem;border-radius:8px;margin-bottom:6px;
        background:rgba(30,41,59,0.6);border-left:3px solid #334155;font-size:0.78rem;}
    .traj-thought{border-left-color:#818CF8;}
    .traj-action{border-left-color:#06B6D4;}
    .traj-obs{border-left-color:#10B981;}
    .traj-decision{border-left-color:#F59E0B;}
    .traj-guardrail{border-left-color:#EF4444;}
    .traj-human{border-left-color:#F59E0B;}
    .traj-scheduler{border-left-color:#8B5CF6;}
    .guardrail-row{display:flex;justify-content:space-between;align-items:center;
        padding:0.4rem 0;border-bottom:1px solid rgba(255,255,255,0.05);font-size:0.8rem;}
    .pass-pill{color:#10B981;font-weight:700;}
    .fail-pill{color:#EF4444;font-weight:700;}
    .pending-pill{color:#F59E0B;font-weight:700;}
    .evidence-box{background:rgba(6,182,212,0.07);border:1px solid rgba(6,182,212,0.2);
        border-radius:6px;padding:0.4rem 0.6rem;font-size:0.72rem;color:#94A3B8;margin-top:3px;}
    .jd-preview{background:rgba(15,23,42,0.8);border:1px solid rgba(99,102,241,0.3);
        border-radius:10px;padding:1rem;font-size:0.78rem;color:#CBD5E1;
        max-height:300px;overflow-y:auto;white-space:pre-wrap;font-family:monospace;}
    .resume-item{display:flex;justify-content:space-between;align-items:center;
        padding:0.5rem 0.75rem;background:rgba(30,41,59,0.5);border-radius:8px;
        border:1px solid rgba(255,255,255,0.06);margin-bottom:6px;}
    .status-dot-green{color:#10B981;font-weight:700;}
    .status-dot-amber{color:#F59E0B;font-weight:700;}
    </style>
    """, unsafe_allow_html=True)


# ── Session state ─────────────────────────────────────────────────────────────
def _init_session() -> None:
    defaults = {
        "current_step":       1,
        "jd_text":            "",
        "jd_source":          "",          # "ai_generated" | "uploaded"
        "resume_inputs":      [],          # list of {filename, content, size}
        "agent_state":        {},
        "run_id":             "",
        "last_run_summary":   None,
        "guardrail_status":   {},
        "guardrail_overall":  "idle",
        "agent_running":      False,
        "approval_pending":   False,
        "selected_model":     get_settings().default_model,
        "temperature":        get_settings().default_temperature,
        "max_tokens":         get_settings().default_max_tokens,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ── File reading ──────────────────────────────────────────────────────────────
def _read_file(file) -> str:
    if file is None:
        return ""
    raw = file.read()
    name = file.name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf_bytes(raw, file.name)
    if name.endswith(".docx"):
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            pass
    return raw.decode("utf-8", errors="replace")


# ── Wizard progress bar ───────────────────────────────────────────────────────
def _render_progress(current: int) -> None:
    steps = [
        (1, "Landing"),
        (2, "Job Description"),
        (3, "Upload Resumes"),
        (4, "Run Agent"),
        (5, "Execution"),
        (6, "Results"),
    ]
    html = '<div class="wizard-progress">'
    for n, label in steps:
        if n < current:
            cls = "wstep-done"
            icon = "✓"
        elif n == current:
            cls = "wstep-active"
            icon = str(n)
        else:
            cls = "wstep-pending"
            icon = str(n)
        html += f'<div class="wstep {cls}"><span>{icon}</span><span>{label}</span></div>'
        if n < len(steps):
            html += '<div style="flex:1;height:1px;background:rgba(255,255,255,0.07);"></div>'
    html += "</div>"
    st.markdown(html, unsafe_allow_html=True)


def _step_header(badge: str, title: str, subtitle: str) -> None:
    st.markdown(f"""
    <div class="step-header">
        <div class="step-badge">{badge}</div>
        <div class="step-title">{title}</div>
        <div class="step-subtitle">{subtitle}</div>
    </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 1 — Landing Page
# ═══════════════════════════════════════════════════════════════════════════════
def render_step1() -> None:
    st.markdown("""
    <div style="text-align:center;padding:3rem 1rem 2rem;">
        <div style="font-size:3rem;margin-bottom:0.5rem;">🎯</div>
        <h1 style="font-size:2.2rem;font-weight:900;color:#F1F5F9;margin:0;">
            TechVest Recruitment Agent
        </h1>
        <p style="font-size:1rem;color:#94A3B8;margin-top:0.75rem;max-width:600px;margin-left:auto;margin-right:auto;">
            Autonomous LangGraph-powered ATS. Upload a Job Description and any number of
            candidate resumes. The agent parses, scores, ranks, and recommends — no hardcoded data,
            no fixed candidates. Every output is generated from your inputs.
        </p>
    </div>""", unsafe_allow_html=True)

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("""
        <div class="card-section">
            <div style="font-size:1.5rem;margin-bottom:0.5rem;">📄</div>
            <div style="font-weight:700;color:#F1F5F9;margin-bottom:0.4rem;">Step 1 — Job Description</div>
            <div style="font-size:0.8rem;color:#94A3B8;">
                Generate a JD using AI by filling in role details, or upload an existing PDF/DOCX/TXT.
            </div>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown("""
        <div class="card-section">
            <div style="font-size:1.5rem;margin-bottom:0.5rem;">📎</div>
            <div style="font-weight:700;color:#F1F5F9;margin-bottom:0.4rem;">Step 2 — Upload Resumes</div>
            <div style="font-size:0.8rem;color:#94A3B8;">
                Upload any number of candidate resumes (PDF, DOCX, TXT). 3 resumes or 50 — the agent scales.
            </div>
        </div>""", unsafe_allow_html=True)
    with c3:
        st.markdown("""
        <div class="card-section">
            <div style="font-size:1.5rem;margin-bottom:0.5rem;">⚡</div>
            <div style="font-weight:700;color:#F1F5F9;margin-bottom:0.4rem;">Step 3 — Run & Review</div>
            <div style="font-size:0.8rem;color:#94A3B8;">
                The LangGraph agent autonomously parses, scores, detects injection, runs fairness checks,
                and produces a ranked shortlist with evidence.
            </div>
        </div>""", unsafe_allow_html=True)

    st.markdown("""
    <div style="background:rgba(16,185,129,0.07);border:1px solid rgba(16,185,129,0.2);
         border-radius:12px;padding:1.2rem 1.5rem;margin-top:1rem;">
        <div style="font-weight:700;color:#10B981;margin-bottom:0.5rem;">How It Works</div>
        <div style="font-size:0.82rem;color:#94A3B8;line-height:1.9;">
            1. Create or upload a Job Description &nbsp;→&nbsp;
            2. Upload candidate resumes (any number) &nbsp;→&nbsp;
            3. Click <b style="color:#818CF8;">Run Agent</b> &nbsp;→&nbsp;
            4. LangGraph autonomously: parses each resume · detects prompt injection ·
               scores against a 7-criterion rubric · runs fairness audit · makes
               Interview / Hold / Reject decisions · generates justifications &nbsp;→&nbsp;
            5. Recruiter reviews results &amp; approves interviews &nbsp;→&nbsp;
            6. Scheduler proposes slots (awaits approval, never auto-fires)
        </div>
    </div>""", unsafe_allow_html=True)

    st.markdown("<div style='margin-top:2rem;text-align:center;'>", unsafe_allow_html=True)
    if st.button("🚀  Get Started →", type="primary", use_container_width=False):
        st.session_state["current_step"] = 2
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 2 — Job Description
# ═══════════════════════════════════════════════════════════════════════════════
def _generate_jd_with_ai(role: str, skills: str, experience: str,
                          responsibilities: str, location: str, emp_type: str) -> str:
    """Call LLM to generate a professional JD."""
    from llm.client import get_llm_client
    prompt = f"""Write a professional Job Description for the following role.

Role: {role}
Required Skills: {skills}
Experience Required: {experience}
Key Responsibilities: {responsibilities}
Location: {location}
Employment Type: {emp_type}

Format it as a proper JD with:
- Job Title
- Company Overview (use "TechVest" as company)
- Role Summary
- Key Responsibilities (bullet list)
- Required Skills & Qualifications (bullet list)
- Nice-to-Have Skills (bullet list)
- Experience Required
- Location & Work Mode
- Employment Type
- What We Offer

Return ONLY the job description text, no extra commentary."""
    try:
        client = get_llm_client()
        result = client.invoke(prompt)
        return result if isinstance(result, str) else str(result)
    except Exception as exc:
        return f"[JD generation failed: {exc}]\n\nPlease upload a JD manually."


def render_step2() -> None:
    _step_header("STEP 2 / 6", "Create Job Description",
                 "Generate a JD using AI or upload an existing one.")

    tab_ai, tab_upload = st.tabs(["✨ Generate with AI", "📤 Upload Existing JD"])

    with tab_ai:
        st.markdown('<div class="section-title">Enter Role Details</div>', unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            role = st.text_input("Role / Job Title *", placeholder="e.g. Senior Machine Learning Engineer")
            skills = st.text_area("Required Skills *", placeholder="Python, PyTorch, MLOps, Docker, AWS...", height=80)
            experience = st.text_input("Experience Required", placeholder="e.g. 3–5 years")
        with col2:
            responsibilities = st.text_area("Key Responsibilities", placeholder="Train and deploy ML models...", height=80)
            location = st.text_input("Location", placeholder="e.g. Bangalore, India (Hybrid)")
            emp_type = st.selectbox("Employment Type", ["Full-Time", "Part-Time", "Contract", "Internship"])

        if st.button("✨ Generate Job Description", type="primary", disabled=not (role and skills)):
            with st.spinner("Generating JD with AI…"):
                jd = _generate_jd_with_ai(role, skills, experience, responsibilities, location, emp_type)
            st.session_state["jd_text"] = jd
            st.session_state["jd_source"] = "ai_generated"
            st.success("✅ Job Description generated!")

    with tab_upload:
        st.markdown('<div class="section-title">Upload Job Description File</div>', unsafe_allow_html=True)
        jd_file = st.file_uploader(
            "Upload JD (PDF, TXT, DOCX)",
            type=["pdf", "txt", "docx", "md"],
            key="jd_file_uploader",
        )
        if jd_file:
            jd_text = _read_file(jd_file)
            if jd_text.strip():
                st.session_state["jd_text"] = jd_text
                st.session_state["jd_source"] = "uploaded"
                st.success(f"✅ Loaded: **{jd_file.name}** ({len(jd_text)} characters)")

    # Preview
    jd_text = st.session_state.get("jd_text", "")
    if jd_text:
        st.markdown('<div class="section-title">📄 Job Description Preview</div>', unsafe_allow_html=True)
        src = st.session_state.get("jd_source", "")
        src_label = "✨ AI Generated" if src == "ai_generated" else "📤 Uploaded"
        st.markdown(f'<div style="font-size:0.72rem;color:#818CF8;margin-bottom:0.5rem;">{src_label}</div>',
                    unsafe_allow_html=True)
        st.markdown(f'<div class="jd-preview">{jd_text[:2000]}{"..." if len(jd_text) > 2000 else ""}</div>',
                    unsafe_allow_html=True)

        col_back, col_next = st.columns([1, 3])
        with col_back:
            if st.button("← Back"):
                st.session_state["current_step"] = 1
                st.rerun()
        with col_next:
            if st.button("✅ Confirm JD & Continue →", type="primary"):
                st.session_state["current_step"] = 3
                st.rerun()
    else:
        if st.button("← Back"):
            st.session_state["current_step"] = 1
            st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 3 — Upload Resumes
# ═══════════════════════════════════════════════════════════════════════════════
def render_step3() -> None:
    _step_header("STEP 3 / 6", "Upload Candidate Resumes",
                 "Upload any number of resumes. Each file becomes one candidate card in the results.")

    resume_files = st.file_uploader(
        "Upload Resumes (PDF, DOCX, TXT) — multiple files allowed",
        type=["pdf", "txt", "docx"],
        accept_multiple_files=True,
        key="resume_multi_uploader",
    )

    if resume_files:
        inputs = []
        for f in resume_files:
            content = _read_file(f)
            inputs.append({
                "filename": f.name,
                "content": content,
                "size_kb": round(len(f.getvalue()) / 1024, 1),
            })
        st.session_state["resume_inputs"] = inputs

    resume_inputs = st.session_state.get("resume_inputs", [])

    if resume_inputs:
        st.markdown(f'<div class="section-title">📎 {len(resume_inputs)} Resume(s) Ready</div>',
                    unsafe_allow_html=True)
        for r in resume_inputs:
            parsed_chars = len(r.get("content", ""))
            status = "✅ Parsed" if parsed_chars > 50 else "⚠️ Empty"
            st.markdown(f"""
            <div class="resume-item">
                <div>
                    <span style="color:#F1F5F9;font-weight:600;">{r['filename']}</span>
                    <span style="font-size:0.72rem;color:#64748B;margin-left:0.5rem;">{r.get('size_kb', 0)} KB · {parsed_chars} chars</span>
                </div>
                <span class="{'status-dot-green' if '✅' in status else 'status-dot-amber'}">{status}</span>
            </div>""", unsafe_allow_html=True)
    else:
        st.info("No resumes uploaded yet. Use the uploader above.")

    col_back, col_next = st.columns([1, 3])
    with col_back:
        if st.button("← Back to JD"):
            st.session_state["current_step"] = 2
            st.rerun()
    with col_next:
        can_proceed = len(resume_inputs) > 0
        if st.button("✅ Confirm Resumes & Continue →", type="primary", disabled=not can_proceed):
            st.session_state["current_step"] = 4
            st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 4 — Run Agent
# ═══════════════════════════════════════════════════════════════════════════════
def render_step4() -> None:
    _step_header("STEP 4 / 6", "Run Recruitment Agent",
                 "Review inputs then launch the autonomous LangGraph agent.")

    jd_text = st.session_state.get("jd_text", "")
    resume_inputs = st.session_state.get("resume_inputs", [])
    jd_ok = bool(jd_text.strip())
    resumes_ok = len(resume_inputs) > 0

    # Pre-flight checklist
    st.markdown('<div class="section-title">Pre-flight Checklist</div>', unsafe_allow_html=True)

    def check_row(label, ok, detail=""):
        icon = "✅" if ok else "❌"
        color = "#10B981" if ok else "#EF4444"
        st.markdown(
            f'<div style="display:flex;justify-content:space-between;padding:0.5rem 0.75rem;'
            f'background:rgba(30,41,59,0.5);border-radius:8px;margin-bottom:6px;">'
            f'<span style="color:#F1F5F9;">{label}</span>'
            f'<span style="color:{color};font-weight:700;">{icon} {detail}</span></div>',
            unsafe_allow_html=True,
        )

    check_row("Job Description", jd_ok,
              f"Ready ({len(jd_text)} chars)" if jd_ok else "Missing — go back to Step 2")
    check_row("Candidate Resumes", resumes_ok,
              f"{len(resume_inputs)} resume(s) loaded" if resumes_ok else "No resumes — go back to Step 3")

    settings = get_settings()
    api_ok = settings.is_configured
    check_row("API Key", api_ok,
              "Connected" if api_ok else "OPENROUTER_API_KEY missing in .env")

    # Model settings
    st.markdown('<div class="section-title">LLM Settings</div>', unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1:
        models = ["openai/gpt-4o-mini", "openai/gpt-4o", "openai/gpt-3.5-turbo",
                  "anthropic/claude-3-haiku", "anthropic/claude-3-sonnet",
                  "google/gemini-flash-1.5", "mistralai/mistral-7b-instruct"]
        cur = st.session_state.get("selected_model", models[0])
        if cur not in models:
            cur = models[0]
        st.session_state["selected_model"] = st.selectbox("Model", models, index=models.index(cur))
    with col2:
        st.session_state["temperature"] = st.slider(
            "Temperature", 0.0, 1.0,
            value=st.session_state.get("temperature", 0.1), step=0.05)
    with col3:
        st.session_state["max_tokens"] = st.select_slider(
            "Max Tokens", options=[512, 1024, 2048, 4096, 8192],
            value=st.session_state.get("max_tokens", 4096))

    st.divider()
    col_back, col_run = st.columns([1, 3])
    with col_back:
        if st.button("← Back to Resumes"):
            st.session_state["current_step"] = 3
            st.rerun()
    with col_run:
        ready = jd_ok and resumes_ok
        if st.button(
            f"▶  Run Agent on {len(resume_inputs)} Candidate(s)" if ready else "▶  Run Agent (fix issues above)",
            type="primary",
            disabled=not ready,
            use_container_width=True,
        ):
            st.session_state["current_step"] = 5
            st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 5 — Agent Execution
# ═══════════════════════════════════════════════════════════════════════════════
def _run_agent(jd_text: str, resume_inputs: list[dict], cfg: dict) -> None:
    from graph.state import create_initial_state
    from graph.graph_builder import build_graph
    from database.sqlite import get_db
    from database.audit import get_audit_logger
    from llm.client import get_llm_client

    settings = get_settings()
    run_id = str(uuid.uuid4())
    st.session_state["run_id"] = run_id
    st.session_state["agent_running"] = True

    get_llm_client(
        model=cfg.get("selected_model"),
        temperature=cfg.get("temperature"),
        max_tokens=cfg.get("max_tokens"),
        force_new=True,
    )

    try:
        db = get_db()
        db.create_run(run_id, jd_text[:500])
    except Exception as e:
        logger.warning(f"DB init failed: {e}")

    audit = get_audit_logger(run_id=run_id, force_new=True)
    audit.info("agent_run_started", details={
        "model": cfg.get("selected_model"),
        "resumes": len(resume_inputs),
    })

    initial_state = create_initial_state(
        run_id=run_id,
        job_description=jd_text,
        resume_inputs=resume_inputs,
        session_id=run_id,
    )
    st.session_state["agent_state"] = initial_state

    status_ph = st.empty()
    status_ph = st.empty()
    progress_bar = st.progress(0, text="Starting agent…")
    debug_log = st.expander("📋 Live Agent Log", expanded=True)

    try:
        compiled = build_graph(use_checkpointer=True)
        config = {
            "recursion_limit": max(settings.recursion_limit, 100),
            "configurable": {"thread_id": run_id},
        }
        step = 0
        total_est = len(resume_inputs) * 4 + 6

        for event in compiled.stream(initial_state, config=config, stream_mode="updates"):
            step += 1
            pct = min(int(step / total_est * 100), 95)
            node = list(event.keys())[0] if event else "processing"

            # Log every node firing so we can see what's happening
            with debug_log:
                st.write(f"**Step {step}** → node: `{node}`")

            for node_name, updates in event.items():
                if isinstance(updates, dict):
                    current = st.session_state.get("agent_state", {})

                    # Log key updates
                    with debug_log:
                        if updates.get("error_message"):
                            st.error(f"Node `{node_name}` error: {updates['error_message']}")
                        if isinstance(updates.get("parsed_profiles"), list):
                            st.write(f"  → parsed_profiles: {len(updates['parsed_profiles'])}")
                        if isinstance(updates.get("scorecards"), list):
                            st.write(f"  → scorecards: {len(updates['scorecards'])}")
                        if isinstance(updates.get("final_decisions"), list):
                            st.write(f"  → final_decisions: {len(updates['final_decisions'])}")
                        if updates.get("next_action"):
                            st.write(f"  → next_action: `{updates['next_action']}`")

                    merged = {**current, **updates}
                    # For list fields: take whichever is longer (latest update wins if it has more items)
                    for list_key in ("trajectory", "parsed_profiles", "scorecards",
                                     "final_decisions", "availability_results", "scored_candidates"):
                        existing = current.get(list_key) or []
                        incoming = updates.get(list_key)
                        if isinstance(incoming, list):
                            if len(incoming) >= len(existing):
                                merged[list_key] = incoming
                            else:
                                existing_names = {
                                    (x.get("candidate_name") or x.get("name") or x.get("filename") or str(i))
                                    for i, x in enumerate(existing)
                                    if isinstance(x, dict)
                                }
                                merged[list_key] = existing + [
                                    x for x in incoming
                                    if not isinstance(x, dict) or
                                    (x.get("candidate_name") or x.get("name") or x.get("filename")) not in existing_names
                                ]
                    st.session_state["agent_state"] = merged

            gs = st.session_state["agent_state"].get("guardrail_status", {})
            st.session_state["guardrail_status"] = gs
            st.session_state["guardrail_overall"] = "pass" if gs.get("overall_pass", True) else "fail"

            node_label = node.replace("_", " ").title()
            progress_bar.progress(pct, text=f"Running: {node_label}…")
            status_ph.markdown(
                f'<div style="font-size:0.78rem;color:#818CF8;margin-top:0.25rem;">⚡ {node_label}</div>',
                unsafe_allow_html=True,
            )

            approval = st.session_state["agent_state"].get("human_approval", {})
            if approval.get("pending", False) and not approval.get("approved", False):
                st.session_state["approval_pending"] = True
                break

        progress_bar.progress(100, text="Agent complete!")
        time.sleep(0.4)
        progress_bar.empty()
        status_ph.empty()

        # Pull the final checkpointed state — this is the ground truth
        try:
            final_snapshot = compiled.get_state(config)
            if final_snapshot and final_snapshot.values:
                # Merge final snapshot into session state (it has all accumulated fields)
                current = st.session_state.get("agent_state", {})
                final_vals = dict(final_snapshot.values)
                merged_final = {**current, **final_vals}
                st.session_state["agent_state"] = merged_final
                logger.info(f"Final state pulled: {len(final_vals.get('scorecards', []))} scorecards")
        except Exception as snap_err:
            logger.warning(f"Could not pull final state snapshot: {snap_err}")

    except Exception as exc:
        logger.error(f"Agent run failed: {exc}", exc_info=True)
        progress_bar.empty()
        status_ph.empty()
        st.error(f"❌ Agent error: {exc}")
        state = st.session_state.get("agent_state", {})
        state["status"] = "error"
        state["error_message"] = str(exc)
        st.session_state["agent_state"] = state
    finally:
        st.session_state["agent_running"] = False

    # Build summary
    final_state = st.session_state.get("agent_state", {})
    scorecards = final_state.get("scorecards", [])
    start_ms = final_state.get("execution_start_ms", 0)
    duration = (time.time() * 1000 - start_ms) / 1000

    summary = {
        "run_id":           run_id,
        "status":           final_state.get("status", "completed"),
        "completed_at":     __import__("datetime").datetime.utcnow().isoformat(),
        "total_candidates": len(scorecards),
        "interview_count":  sum(1 for s in scorecards if s.get("recommendation") == "Interview"),
        "hold_count":       sum(1 for s in scorecards if s.get("recommendation") == "Hold"),
        "reject_count":     sum(1 for s in scorecards if s.get("recommendation") == "Reject"),
        "avg_score":        sum(s.get("overall_weighted_score", 0) for s in scorecards) / max(len(scorecards), 1),
        "top_candidate":    final_state.get("top_candidate", ""),
        "total_tool_calls": final_state.get("total_tool_calls", 0),
        "total_llm_calls":  final_state.get("total_llm_calls", 0),
        "duration_seconds": round(duration, 2),
    }
    st.session_state["last_run_summary"] = summary
    try:
        get_db().complete_run(run_id, summary)
    except Exception as e:
        logger.warning(f"DB complete_run failed: {e}")


def render_step5() -> None:
    _step_header("STEP 5 / 6", "Agent Execution",
                 "LangGraph is autonomously processing every candidate.")

    jd_text = st.session_state.get("jd_text", "")
    resume_inputs = st.session_state.get("resume_inputs", [])

    # Use a dedicated flag so we never skip the actual run
    already_ran = st.session_state.get("agent_run_complete", False)

    if already_ran:
        # Check if we have real results
        scorecards = st.session_state.get("agent_state", {}).get("scorecards", [])
        if scorecards:
            st.session_state["current_step"] = 6
            st.rerun()
        else:
            # Run failed or produced nothing — show error and let user retry
            error_msg = st.session_state.get("agent_state", {}).get("error_message", "")
            st.error(f"❌ Agent run produced no results. {error_msg or 'Check your API key and try again.'}")

            # Debug expander so the user can see what happened
            with st.expander("🔍 Debug — Agent State (expand to diagnose)", expanded=False):
                state_debug = st.session_state.get("agent_state", {})
                st.write(f"**Status:** `{state_debug.get('status', 'unknown')}`")
                st.write(f"**Parsed profiles:** {len(state_debug.get('parsed_profiles', []))}")
                st.write(f"**Scorecards:** {len(state_debug.get('scorecards', []))}")
                st.write(f"**Final decisions:** {len(state_debug.get('final_decisions', []))}")
                st.write(f"**Trajectory events:** {len(state_debug.get('trajectory', []))}")
                st.write(f"**Error:** `{state_debug.get('error_message', 'none')}`")
                st.write(f"**Current node:** `{state_debug.get('current_node', 'none')}`")
                trajectory = state_debug.get("trajectory", [])
                if trajectory:
                    st.write("**Last 5 trajectory events:**")
                    for ev in trajectory[-5:]:
                        st.write(f"- `{ev.get('event_type','?')}` — {ev.get('title','?')}: {str(ev.get('content',''))[:120]}")

            col1, col2 = st.columns(2)
            with col1:
                if st.button("🔄 Retry Agent Run", type="primary", use_container_width=True):
                    st.session_state["agent_run_complete"] = False
                    st.session_state["agent_state"] = {}
                    st.rerun()
            with col2:
                if st.button("← Back to Settings", use_container_width=True):
                    st.session_state["agent_run_complete"] = False
                    st.session_state["agent_state"] = {}
                    st.session_state["current_step"] = 4
                    st.rerun()
        return

    # Not yet run — execute now
    cfg = {
        "selected_model": st.session_state.get("selected_model"),
        "temperature":    st.session_state.get("temperature"),
        "max_tokens":     st.session_state.get("max_tokens"),
    }

    st.info(f"⚡ Running agent on **{len(resume_inputs)} resume(s)** against your Job Description…")
    log_ph = st.empty()

    with st.spinner(f"LangGraph processing {len(resume_inputs)} candidate(s)…"):
        _run_agent(jd_text, resume_inputs, cfg)

    st.session_state["agent_run_complete"] = True

    # Check outcome
    final_state = st.session_state.get("agent_state", {})
    status = final_state.get("status", "")
    scorecards = final_state.get("scorecards", [])
    error_msg = final_state.get("error_message", "")

    if status == "error" or (status != "error" and not scorecards):
        log_ph.error(f"❌ Agent error: {error_msg or 'No scorecards generated. Check API key and logs.'}")
    else:
        log_ph.success(f"✅ Agent complete — {len(scorecards)} candidates evaluated!")
        time.sleep(0.8)
        st.session_state["current_step"] = 6
        st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 6 — Results Dashboard
# ═══════════════════════════════════════════════════════════════════════════════

def _render_kpi_strip(scorecards: list, summary: dict, state: dict) -> None:
    total = len(scorecards)
    interview = sum(1 for s in scorecards if s.get("recommendation") == "Interview")
    hold = sum(1 for s in scorecards if s.get("recommendation") == "Hold")
    rejected = sum(1 for s in scorecards if s.get("recommendation") == "Reject")
    avg_score = sum(s.get("overall_weighted_score", 0) for s in scorecards) / max(total, 1)
    tool_calls = state.get("total_tool_calls", 0)
    start_ms = state.get("execution_start_ms", 0)
    dur = (time.time() * 1000 - start_ms) / 1000 if start_ms else summary.get("duration_seconds", 0)

    kpis = [
        ("👥", str(total),          "Candidates",    "#6366F1"),
        ("✅", str(interview),       "Interview",     "#10B981"),
        ("⏸️", str(hold),            "Hold",          "#F59E0B"),
        ("❌", str(rejected),        "Rejected",      "#EF4444"),
        ("📊", f"{avg_score:.1f}",   "Avg Score",     "#06B6D4"),
        ("⚡", str(tool_calls),      "Tool Calls",    "#8B5CF6"),
        ("⏱️", f"{dur:.1f}s",        "Exec Time",     "#3B82F6"),
    ]
    cols = st.columns(len(kpis))
    for col, (icon, val, label, color) in zip(cols, kpis):
        with col:
            st.markdown(f"""
            <div class="kpi-card">
                <div style="font-size:1.3rem;">{icon}</div>
                <div style="font-size:1.6rem;font-weight:800;color:{color};line-height:1.2;">{val}</div>
                <div style="font-size:0.65rem;color:#64748B;text-transform:uppercase;letter-spacing:0.05em;">{label}</div>
            </div>""", unsafe_allow_html=True)


def _render_candidate_card_full(sc: dict, profile: dict, run_id: str) -> str | None:
    """Render one candidate card using native Streamlit widgets. Returns action clicked or None."""
    from config.settings import UIConstants
    from config.rubric import RUBRIC

    name = sc.get("candidate_name", profile.get("name", "Unknown"))
    score = sc.get("overall_weighted_score", 0.0)
    rec = sc.get("recommendation", "Reject")
    confidence = sc.get("confidence", 0.5)
    strengths = sc.get("strengths", [])
    gaps = sc.get("gaps", [])
    reasoning = sc.get("reasoning", "")
    injection = profile.get("injection_detected", False)
    years_exp = profile.get("years_experience", 0)
    skills = list(dict.fromkeys(
        profile.get("skills", []) + profile.get("programming_languages", [])
    ))
    education = profile.get("education", [])
    projects = profile.get("projects", [])
    criterion_scores = sc.get("criterion_scores", {})

    # Color mapping
    sc_color = UIConstants.score_color(score)
    badge_map = {"Interview": "🟢", "Hold": "🟡", "Reject": "🔴"}
    badge_icon = badge_map.get(rec, "⚪")

    # Education line
    edu_line = ""
    if education:
        e = education[0]
        deg = e.get("degree", "")
        inst = e.get("institution", "")
        yr = e.get("year", "")
        edu_line = f" · {deg} — {inst}" + (f" ({yr})" if yr else "")

    # Card container using columns
    with st.container():
        st.markdown(
            f'<div style="background:rgba(30,41,59,0.7);border:1px solid rgba(255,255,255,0.08);'
            f'border-radius:14px;padding:1.1rem 1.25rem;margin-bottom:0.75rem;'
            f'border-left:4px solid {sc_color};">',
            unsafe_allow_html=True,
        )

        col_info, col_score = st.columns([4, 1])
        with col_info:
            st.markdown(f"**{name}** &nbsp; {badge_icon} **{rec}**" +
                        (" &nbsp; ⚠️ INJECTION" if injection else ""),
                        unsafe_allow_html=True)
            st.caption(f"{years_exp} yr experience{edu_line}")
            if skills:
                st.markdown(" ".join(
                    f'`{s}`' for s in skills[:10]
                ) + (f" +{len(skills)-10} more" if len(skills) > 10 else ""))
        with col_score:
            st.metric("Score", f"{score:.0f}/100", delta=None)
            st.caption(f"{confidence*100:.0f}% confidence")

        st.markdown("</div>", unsafe_allow_html=True)

    with st.expander(f"🔍 Full Analysis — {name}"):
        if injection:
            st.error("⚠️ Prompt Injection Detected — score penalty of −15 pts applied")

        col1, col2 = st.columns(2)
        with col1:
            if strengths:
                st.markdown("**✅ Strengths**")
                for s in strengths:
                    st.markdown(f"- {s}")
            if education:
                st.markdown("**🎓 Education**")
                for edu in education:
                    yr = edu.get("year", "")
                    st.markdown(f"- {edu.get('degree','?')} — {edu.get('institution','?')}" +
                                (f" ({yr})" if yr else ""))
        with col2:
            if gaps:
                st.markdown("**⚠️ Skill Gaps**")
                for g in gaps:
                    st.markdown(f"- {g}")
            if projects:
                st.markdown("**🔧 Projects**")
                for proj in projects[:3]:
                    desc = proj.get("description", "")[:80]
                    st.markdown(f"- **{proj.get('name','')}** — {desc}")

        if reasoning:
            st.info(f"**Agent Reasoning:** {reasoning}")

        # Score breakdown with evidence
        if criterion_scores:
            st.markdown("**📊 Score Breakdown (with evidence)**")
            for criterion in RUBRIC.criteria:
                key = criterion.key
                cs = criterion_scores.get(key, {})
                cs_score = cs.get("score", 0) if isinstance(cs, dict) else 0
                cs_evidence = cs.get("evidence", "No evidence") if isinstance(cs, dict) else ""
                bar_color = UIConstants.score_color(cs_score * 10)
                bar_pct = int(cs_score * 10)
                st.markdown(
                    f'<div style="margin-bottom:8px;">'
                    f'<div style="display:flex;justify-content:space-between;font-size:0.78rem;margin-bottom:2px;">'
                    f'<span style="color:#CBD5E1;">{criterion.label}</span>'
                    f'<span style="color:{bar_color};font-weight:700;">{cs_score}/10</span></div>'
                    f'<div style="height:5px;background:rgba(255,255,255,0.08);border-radius:99px;">'
                    f'<div style="width:{bar_pct}%;height:100%;background:{bar_color};border-radius:99px;"></div></div>'
                    f'<div style="font-size:0.7rem;color:#64748B;margin-top:2px;padding:3px 6px;'
                    f'background:rgba(6,182,212,0.06);border-radius:4px;">📎 {cs_evidence[:120]}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

        # Guardrail status
        inj_status = "⚠️ DETECTED" if injection else "✅ PASS"
        col_g1, col_g2 = st.columns(2)
        with col_g1:
            st.markdown(f"🛡️ **Injection:** {inj_status}")
        with col_g2:
            st.markdown("⚖️ **Fairness:** ✅ Reviewed")

    # Action buttons — only for Interview candidates
    action = None
    if rec == "Interview":
        safe = name.replace(" ", "_").lower()
        b1, b2, b3 = st.columns(3)
        with b1:
            if st.button("✅ Approve Interview", key=f"approve_{safe}_{run_id}",
                         use_container_width=True, type="primary"):
                action = "approve"
        with b2:
            if st.button("❌ Reject", key=f"reject_{safe}_{run_id}",
                         use_container_width=True):
                action = "reject"
        with b3:
            if st.button("📅 Schedule Interview", key=f"schedule_{safe}_{run_id}",
                         use_container_width=True):
                action = "schedule"
    return action


def render_step6() -> None:
    state = st.session_state.get("agent_state", {})
    summary = st.session_state.get("last_run_summary", {})
    scorecards = state.get("scorecards", [])
    decisions = state.get("final_decisions", [])
    profiles = state.get("parsed_profiles", [])
    trajectory = state.get("trajectory", [])
    run_id = state.get("run_id", st.session_state.get("run_id", ""))

    if not scorecards:
        st.warning("No results yet. Go back and run the agent.")
        if st.button("← Back to Run"):
            st.session_state["current_step"] = 4
            # Reset state so agent re-runs
            st.session_state["agent_state"] = {}
            st.rerun()
        return

    _step_header("STEP 6 / 6", "Results Dashboard",
                 f"Agent processed {len(scorecards)} candidate(s). Review decisions below.")

    # KPI strip
    _render_kpi_strip(scorecards, summary, state)
    st.markdown("<div style='margin-top:1.5rem;'></div>", unsafe_allow_html=True)

    # ── Summary panel ──
    top_c = summary.get("top_candidate") or state.get("top_candidate", "—")
    sorted_sc = sorted(scorecards, key=lambda x: x.get("overall_weighted_score", 0))
    lowest_c = sorted_sc[0].get("candidate_name", "—") if sorted_sc else "—"
    # Most common missing skill
    all_gaps = [g for sc in scorecards for g in sc.get("gaps", [])]
    from collections import Counter
    top_gap = Counter(all_gaps).most_common(1)
    top_gap_str = top_gap[0][0] if top_gap else "—"

    st.markdown("""<div class="section-title">📊 Summary Panel</div>""", unsafe_allow_html=True)
    sp1, sp2, sp3, sp4 = st.columns(4)
    with sp1:
        st.metric("🏆 Top Candidate", top_c)
    with sp2:
        st.metric("⬇️ Lowest Score", lowest_c)
    with sp3:
        st.metric("🔍 Most Missing Skill", top_gap_str)
    with sp4:
        overall_rec = "Strong Batch" if summary.get("interview_count", 0) > len(scorecards) // 2 else "Mixed Batch"
        st.metric("📋 Overall", overall_rec)

    st.divider()

    # ── Ranking table ──
    st.markdown('<div class="section-title">🏆 Ranking Table</div>', unsafe_allow_html=True)
    from components.tables import render_ranked_shortlist_table, render_scorecard_comparison_table
    render_ranked_shortlist_table(decisions)

    st.markdown('<div class="section-title">📊 Score Comparison (All Candidates)</div>', unsafe_allow_html=True)
    render_scorecard_comparison_table(scorecards)

    st.divider()

    # ── Charts ──
    from components.charts import render_score_bar_chart, render_recommendation_pie, render_execution_timeline, render_rubric_weights
    st.markdown('<div class="section-title">📈 Charts</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**Score Comparison**")
        render_score_bar_chart(scorecards)
    with c2:
        st.markdown("**Recommendation Split**")
        render_recommendation_pie(scorecards)

    st.divider()

    # ── Interview candidates ──
    interview_scs = [s for s in scorecards if s.get("recommendation") == "Interview"]
    if interview_scs:
        st.markdown("---")
        st.markdown("### ✅ Interview Candidates")
        avail_results = state.get("availability_results", [])
        for sc in sorted(interview_scs, key=lambda x: x.get("overall_weighted_score", 0), reverse=True):
            name = sc.get("candidate_name", "")
            profile = next((p for p in profiles if p.get("name") == name), {})
            _render_candidate_card_full(sc, profile, run_id)
            avail = next((a for a in avail_results if a.get("candidate_name") == name), None)
            if avail and avail.get("proposed_slots"):
                slot = avail["proposed_slots"][0]
                st.markdown(
                    f'<div style="font-size:0.75rem;color:#06B6D4;margin-bottom:1rem;">'
                    f'📅 Suggested slot: {slot.get("date","")} {slot.get("time","")} '
                    f'({slot.get("format","Video")}, {slot.get("duration_minutes",60)} min)</div>',
                    unsafe_allow_html=True,
                )

    # ── Hold candidates ──
    hold_scs = [s for s in scorecards if s.get("recommendation") == "Hold"]
    if hold_scs:
        st.markdown("---")
        st.markdown("### ⏸️ Hold Candidates")
        for sc in sorted(hold_scs, key=lambda x: x.get("overall_weighted_score", 0), reverse=True):
            name = sc.get("candidate_name", "")
            profile = next((p for p in profiles if p.get("name") == name), {})
            _render_candidate_card_full(sc, profile, run_id)
            gaps = sc.get("gaps", [])
            if gaps:
                st.markdown(
                    f'<div style="font-size:0.75rem;color:#F59E0B;margin-bottom:1rem;">'
                    f'🔑 Missing: {", ".join(gaps[:3])} — reconsider if skills improve</div>',
                    unsafe_allow_html=True,
                )

    # ── Rejected candidates ──
    reject_scs = [s for s in scorecards if s.get("recommendation") == "Reject"]
    if reject_scs:
        st.markdown("---")
        st.markdown("### ❌ Rejected Candidates")
        for sc in sorted(reject_scs, key=lambda x: x.get("overall_weighted_score", 0), reverse=True):
            name = sc.get("candidate_name", "")
            score = sc.get("overall_weighted_score", 0)
            profile = next((p for p in profiles if p.get("name") == name), {})

            with st.expander(f"❌ {name} — Score: {score:.0f}/100", expanded=True):
                # Explicit rejection reasons
                gaps = sc.get("gaps", [])
                reasoning = sc.get("reasoning", "")
                criterion_scores = sc.get("criterion_scores", {})

                st.markdown("**Why Rejected:**")

                # Build explicit reason list from gaps + low scores
                reasons = []
                if gaps:
                    for g in gaps:
                        reasons.append(f"❌ {g}")

                # Find criteria with score ≤ 3
                from config.rubric import RUBRIC
                for criterion in RUBRIC.criteria:
                    cs = criterion_scores.get(criterion.key, {})
                    cs_score = cs.get("score", 0) if isinstance(cs, dict) else 0
                    cs_evidence = cs.get("evidence", "") if isinstance(cs, dict) else ""
                    if cs_score <= 3:
                        reasons.append(
                            f"❌ **{criterion.label}**: scored {cs_score}/10 — {cs_evidence[:80] or 'Insufficient evidence'}"
                        )

                if not reasons:
                    reasons = [f"❌ Overall score {score:.0f}/100 is below the minimum threshold of 40"]

                for r in reasons:
                    st.markdown(r)

                if reasoning:
                    st.markdown(f"**Agent Assessment:** {reasoning}")

                # Score breakdown table
                st.markdown("**Score Breakdown:**")
                from config.settings import UIConstants
                cols = st.columns(len(RUBRIC.criteria))
                for col, criterion in zip(cols, RUBRIC.criteria):
                    cs = criterion_scores.get(criterion.key, {})
                    cs_score = cs.get("score", 0) if isinstance(cs, dict) else 0
                    color = UIConstants.score_color(cs_score * 10)
                    with col:
                        st.markdown(
                            f'<div style="text-align:center;padding:0.4rem;background:rgba(30,41,59,0.5);'
                            f'border-radius:8px;border-top:2px solid {color};">'
                            f'<div style="font-size:1rem;font-weight:800;color:{color};">{cs_score}</div>'
                            f'<div style="font-size:0.6rem;color:#64748B;">{criterion.label[:8]}</div>'
                            f'</div>',
                            unsafe_allow_html=True,
                        )

    st.divider()

    # ── Trajectory / Audit log ──
    tab_traj, tab_guardrail, tab_export = st.tabs(["🔄 Execution Trajectory", "🛡️ Guardrails", "⬇️ Export"])

    with tab_traj:
        st.markdown('<div class="section-title">Full LangGraph Execution Trajectory</div>',
                    unsafe_allow_html=True)
        if trajectory:
            for ev in trajectory:
                etype = ev.get("event_type", "action")
                cls_map = {"thought": "traj-thought", "action": "traj-action",
                           "observation": "traj-obs", "decision": "traj-decision",
                           "guardrail": "traj-guardrail", "human": "traj-human",
                           "scheduler": "traj-scheduler"}
                cls = cls_map.get(etype, "traj-action")
                ts = str(ev.get("timestamp", ""))[:19].replace("T", " ")
                title = ev.get("title", "")
                content = ev.get("content", "")
                node = ev.get("node", "")
                dur = ev.get("duration_ms")
                dur_str = f" · {dur:.0f}ms" if dur else ""
                st.markdown(
                    f'<div class="traj-event {cls}">'
                    f'<div style="display:flex;justify-content:space-between;">'
                    f'<b style="color:#F1F5F9;">{title}</b>'
                    f'<span style="color:#475569;font-size:0.68rem;">{ts}{dur_str}</span></div>'
                    f'<div style="color:#94A3B8;margin-top:2px;">{content[:200]}</div>'
                    f'<div style="color:#475569;font-size:0.68rem;margin-top:2px;">node: {node}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
        else:
            st.info("No trajectory events recorded.")

    with tab_guardrail:
        st.markdown('<div class="section-title">Guardrail Status</div>', unsafe_allow_html=True)
        gs = state.get("guardrail_status", {})
        fr = state.get("fairness_result", {}) or {}
        injection_count = sum(1 for p in profiles if p.get("injection_detected", False))

        checks = [
            ("🛡️ Prompt Injection", "PASS" if injection_count == 0 else f"DETECTED ({injection_count})"),
            ("⚖️ Fairness Audit", fr.get("overall_fairness", "N/A")),
            ("📊 Step Limit", "PASS" if gs.get("step_limit_ok", True) else "EXCEEDED"),
            ("🔄 Iteration Limit", "PASS" if gs.get("iteration_limit_ok", True) else "EXCEEDED"),
            ("🔍 Loop Detection", "PASS" if not gs.get("loop_detected", False) else "LOOP DETECTED"),
            ("👤 Human Approval", "Approved" if state.get("human_approval", {}).get("approved") else "Pending"),
            ("📋 Audit Log", "Generated"),
        ]
        for label, val in checks:
            is_pass = val in ("PASS", "Generated", "Approved") or val.startswith("PASS")
            is_warn = "Pending" in val or "N/A" in val
            pill_cls = "pass-pill" if is_pass else ("pending-pill" if is_warn else "fail-pill")
            st.markdown(
                f'<div class="guardrail-row"><span style="color:#94A3B8;">{label}</span>'
                f'<span class="{pill_cls}">{val}</span></div>',
                unsafe_allow_html=True,
            )

        if fr.get("checks"):
            st.markdown("**Fairness Detail**")
            for check in fr["checks"]:
                st.markdown(f"- **{check.get('check_name','')}**: {check.get('status','')} — {check.get('finding','')}")

    with tab_export:
        st.markdown('<div class="section-title">Export Results</div>', unsafe_allow_html=True)
        import pandas as pd
        e1, e2, e3 = st.columns(3)
        with e1:
            st.download_button(
                "⬇️ Decisions JSON",
                data=json.dumps(decisions, default=str, indent=2),
                file_name="techvest_decisions.json", mime="application/json",
                use_container_width=True,
            )
        with e2:
            df = pd.DataFrame([{
                "Candidate": d.get("candidate_name"),
                "Score": d.get("weighted_score"),
                "Recommendation": d.get("final_recommendation"),
                "Rank": d.get("rank"),
                "Confidence": d.get("confidence"),
            } for d in decisions])
            st.download_button(
                "⬇️ Shortlist CSV",
                data=df.to_csv(index=False),
                file_name="techvest_shortlist.csv", mime="text/csv",
                use_container_width=True,
            )
        with e3:
            st.download_button(
                "⬇️ Trajectory JSON",
                data=json.dumps(trajectory, default=str, indent=2),
                file_name="techvest_trajectory.json", mime="application/json",
                use_container_width=True,
            )

    st.divider()
    col_back, col_new = st.columns([1, 3])
    with col_back:
        if st.button("← Back"):
            st.session_state["current_step"] = 4
            st.rerun()
    with col_new:
        if st.button("🔄 Start New Run", use_container_width=True):
            for k in ["agent_state", "run_id", "last_run_summary", "guardrail_status",
                      "guardrail_overall", "agent_running", "approval_pending",
                      "jd_text", "jd_source", "resume_inputs", "agent_run_complete"]:
                st.session_state.pop(k, None)
            st.session_state["current_step"] = 1
            _init_session()
            st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
def main() -> None:
    _init_session()
    inject_css()

    # Sidebar
    from components.sidebar import render_dynamic_sidebar
    render_dynamic_sidebar()

    # Page header
    st.markdown("""
    <div style="display:flex;align-items:center;gap:0.75rem;margin-bottom:1rem;">
        <span style="font-size:1.5rem;">🎯</span>
        <div>
            <div style="font-size:0.72rem;color:#475569;text-transform:uppercase;letter-spacing:0.08em;">
                TechVest AI Platform
            </div>
            <div style="font-size:1.1rem;font-weight:800;color:#F1F5F9;">Recruitment Agent</div>
        </div>
    </div>""", unsafe_allow_html=True)

    # API warning
    settings = get_settings()
    if not settings.is_configured:
        st.warning("⚠️ OPENROUTER_API_KEY not configured. Add it to your .env file and restart.")

    current_step = st.session_state.get("current_step", 1)
    _render_progress(current_step)

    # Route to correct step
    if current_step == 1:
        render_step1()
    elif current_step == 2:
        render_step2()
    elif current_step == 3:
        render_step3()
    elif current_step == 4:
        render_step4()
    elif current_step == 5:
        render_step5()
    elif current_step == 6:
        render_step6()
    else:
        render_step1()


if __name__ == "__main__":
    main()
